"""从仓库 docs/rag 等目录加载文本，按模块切块供 RAG 使用（启动时缓存）。

知识库分区：
- policy / subsidy：docs/rag/subsidy（Standard_Subsidy_Policy）
- soul_window / mental_health：docs/rag/mental_health（Mental_Health_Protocol）
"""

from __future__ import annotations

import logging
import os
import re
from functools import lru_cache
from pathlib import Path

logger = logging.getLogger(__name__)

RAG_SCOPES = frozenset({"policy", "soul_window"})


def _read_text_file(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def _read_docx(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    parts: list[str] = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _chunk_text(text: str, source_label: str, *, max_len: int = 1800, overlap: int = 200) -> list[tuple[str, str]]:
    t = text.strip()
    if not t:
        return []
    if len(t) <= max_len:
        return [(source_label, t)]
    out: list[tuple[str, str]] = []
    start = 0
    while start < len(t):
        end = min(start + max_len, len(t))
        out.append((source_label, t[start:end]))
        if end >= len(t):
            break
        start = max(0, end - overlap)
    return out


def _repo_root() -> Path:
    here = Path(__file__).resolve().parent
    return here.parent.parent


def _discover_doc_roots(scope: str = "policy") -> list[Path]:
    """按对话模块发现文档根目录；资助与心理分库，避免混检。"""
    if scope not in RAG_SCOPES:
        scope = "policy"

    repo = _repo_root()
    backend = Path(__file__).resolve().parent.parent
    roots: list[Path] = []

    rag_root = repo / "docs" / "rag"
    if rag_root.is_dir():
        if scope == "policy":
            sub = rag_root / "subsidy"
            if sub.is_dir():
                roots.append(sub)
        else:
            sub = rag_root / "mental_health"
            if sub.is_dir():
                roots.append(sub)

    if scope == "policy":
        # Docker
        if (backend / "docs" / "rag" / "subsidy").is_dir():
            roots.append(backend / "docs" / "rag" / "subsidy")
        if (backend / "public_docs").is_dir():
            roots.append(backend / "public_docs")
        if (repo / "frontend" / "public" / "docs").is_dir():
            roots.append(repo / "frontend" / "public" / "docs")

    for raw in os.environ.get("RAG_EXTRA_DIRS", "").split(","):
        raw = raw.strip()
        if not raw:
            continue
        p = Path(raw).expanduser()
        if p.is_dir():
            roots.append(p)

    seen: set[str] = set()
    out: list[Path] = []
    for r in roots:
        key = str(r.resolve())
        if key not in seen:
            seen.add(key)
            out.append(r)
    return out


def _should_skip_file(rel_posix: str, scope: str) -> bool:
    """跳过演示/非政策类 Markdown。"""
    name = rel_posix.lower()
    if scope == "policy":
        skip_markers = ("aigc-app-ppt", "network-communication-ppt")
        if any(m in name for m in skip_markers):
            return True
    return False


def _load_raw_chunks(scope: str = "policy") -> tuple[list[str], list[str]]:
    """返回 (chunks, source_labels)，一一对应。"""
    chunks: list[str] = []
    sources: list[str] = []
    for root in _discover_doc_roots(scope):
        try:
            paths = sorted(root.rglob("*"))
        except OSError as e:
            logger.warning("RAG 无法遍历 %s: %s", root, e)
            continue
        for path in paths:
            if not path.is_file():
                continue
            suf = path.suffix.lower()
            if suf not in (".md", ".docx", ".txt"):
                continue
            try:
                rel = path.relative_to(root)
            except ValueError:
                continue
            if _should_skip_file(rel.as_posix(), scope):
                continue
            label = f"{root.name}/{rel.as_posix()}"
            try:
                if suf == ".docx":
                    text = _read_docx(path)
                else:
                    text = _read_text_file(path)
            except Exception:
                logger.exception("RAG 跳过文件（读取失败）: %s", path)
                continue
            for src, ch in _chunk_text(text, label):
                sources.append(src)
                chunks.append(ch)
    if chunks:
        logger.info(
            "RAG[%s] 已加载 %s 个文本块，来自 %s",
            scope,
            len(chunks),
            [str(r) for r in _discover_doc_roots(scope)],
        )
    else:
        logger.info("RAG[%s] 未加载任何文档块", scope)
    return chunks, sources


@lru_cache(maxsize=4)
def get_rag_chunks_and_sources(scope: str = "policy") -> tuple[tuple[str, ...], tuple[str, ...]]:
    """进程内缓存；修改文档后需重启后端生效。scope: policy | soul_window"""
    if scope not in RAG_SCOPES:
        scope = "policy"
    c, s = _load_raw_chunks(scope)
    return tuple(c), tuple(s)


def clear_rag_cache() -> None:
    get_rag_chunks_and_sources.cache_clear()
