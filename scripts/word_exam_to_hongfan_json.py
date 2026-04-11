"""
将「考研政治」真题目录下的 Word（.docx）转为红帆知海题库 JSON。

默认扫描：项目根下「03 2009-2025真题和解析」内的「2009-2025政治真题」子目录（递归 *.docx）。
「2009-2025政治解析」目录当前多为 PDF，本脚本不解析 PDF；答案字段为占位说明，可后续手工或扩展 PDF 导入。

用法（在项目根）:
  backend\\.venv\\Scripts\\python.exe scripts/word_exam_to_hongfan_json.py
  backend\\.venv\\Scripts\\python.exe scripts/word_exam_to_hongfan_json.py --out 01历年真题及解析/hongfan-bank.json

依赖: pip install python-docx
"""
from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path

try:
    from docx import Document
except ImportError as e:  # pragma: no cover
    print("请先安装: pip install python-docx", file=sys.stderr)
    raise SystemExit(1) from e

# 题号行：10–38 与 1–9（1–9 时点后不能紧跟数字，避免「5. 4万亿美元」「9. 6%」断行）
_NEXT_Q_10_38 = re.compile(r"^((?:3[0-8]|[12]\d))[.．]\s*(.*)$")
_NEXT_Q_1_9 = re.compile(r"^([1-9])[.．]\s*(.*)$")


def try_question_start(line: str) -> tuple[int, str] | None:
    """若本行以有效题号开头，返回 (题号, 余下题干)；否则 None。"""
    m2 = _NEXT_Q_10_38.match(line)
    if m2:
        return int(m2.group(1)), (m2.group(2) or "").strip()
    m1 = _NEXT_Q_1_9.match(line)
    if not m1:
        return None
    q = int(m1.group(1))
    tail = (m1.group(2) or "").strip()
    # 排除「5. 4万亿美元」「9. 6%」等小数/统计断行；保留「1. 1921年中国」类以年份开头
    if tail and tail[0].isdigit():
        if re.match(r"^\d{4}年", tail):
            return q, tail
        return None
    return q, tail
# 大题标题（新一节），题块在此结束
SECTION = re.compile(r"^[一二三四五六七八九十百千]+[、．.]")
# 行内粘连：前文汉字/标点 + 题号 + 点，后接题干引号或汉字（避免拆 2008年5月 等）
GLUE_BEFORE_Q = re.compile(
    r"(?<=[\u4e00-\u9fff\u3002\u3001\)」”])\s*"
    r"([1-9]\d{0,2})([.．])(?=[\s\u200b\u201c\u2018\u201d\u300a\u300b【「])"
)
# 选项行：行首 A. / A．
OPTION_PREFIX = re.compile(r"^[ABCDabcd][\.．、]\s*")
# 同一行内出现 A.…B.…C.…D.…（允许跨行由调用方拼接）
ABCD_ONE_LINE = re.compile(
    r"A[.．].{0,600}?B[.．].{0,600}?C[.．].{0,600}?D[.．]", re.DOTALL
)
SECTION_CHOICE = re.compile(r"^[一二三四五六七八九十]+[、．.].*选择题")
SECTION_ANALYSIS = re.compile(
    r"^[一二三四五六七八九十]+[、．.].*(?:分析题|材料分析题)"
)
FOOTER_PAGE = re.compile(r"考研政治.*第\s*\d+\s*页|共\s*\d+\s*页")


def find_politics_zhen_dir(root: Path) -> Path | None:
    """定位 …/03 …/2009-2025政治真题。"""
    for p in root.rglob("*2009-2025政治真题*"):
        if p.is_dir():
            return p
    return None


def split_glued_numbering(lines: list[str]) -> list[str]:
    """把同一行末尾粘连的下一题题号拆成新行。"""
    out: list[str] = []
    for line in lines:
        if GLUE_BEFORE_Q.search(line):
            line = GLUE_BEFORE_Q.sub(r"\n\1\2", line)
        for piece in line.split("\n"):
            t = piece.strip()
            if t:
                out.append(t)
    return out


def paragraphs_to_lines(doc: Document) -> list[str]:
    lines: list[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").replace("\u00a0", " ").strip()
        if t:
            lines.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = " | ".join((c.text or "").replace("\u00a0", " ").strip() for c in row.cells)
            cells = cells.strip()
            if cells:
                lines.append(cells)
    return split_glued_numbering(lines)


def extract_year(filename: str) -> str:
    m = re.search(r"(20\d{2})", filename)
    return m.group(1) if m else "unknown"


def parse_numbered_blocks(lines: list[str]) -> list[tuple[int, str]]:
    """按行首题号切分；题干与选项一并写入同一文本块。"""
    blocks: list[tuple[int, str]] = []
    i = 0
    n = len(lines)
    while i < n:
        head = try_question_start(lines[i])
        if not head:
            i += 1
            continue
        qn, tail = head
        parts: list[str] = []
        if tail:
            parts.append(tail)
        i += 1
        while i < n:
            nxt = lines[i]
            if try_question_start(nxt):
                break
            if SECTION.match(nxt):
                break
            parts.append(nxt)
            i += 1
        text = "\n".join(parts).strip()
        if text:
            blocks.append((qn, text))
    return blocks


def _is_noise_line(s: str) -> bool:
    if FOOTER_PAGE.search(s):
        return True
    if "研究生入学考试" in s and "试题" in s:
        return True
    if re.match(r"^绝密", s):
        return True
    return False


def _find_section_starts(lines: list[str]) -> dict[str, int]:
    """返回 '一'/'二'/'三' 对应大题说明行下标（含该行）。"""
    idx: dict[str, int] = {}
    for i, ln in enumerate(lines):
        if "一" in ln[:3] and SECTION_CHOICE.match(ln) and "单项" in ln:
            idx.setdefault("一", i)
        if "二" in ln[:3] and SECTION_CHOICE.match(ln) and "多项" in ln:
            idx.setdefault("二", i)
        if SECTION_ANALYSIS.match(ln) or ("三" in ln[:3] and "分析" in ln):
            idx.setdefault("三", i)
    return idx


def _consume_options_block(lines: list[str], i: int, end: int) -> tuple[str, int]:
    """从 lines[i] 起读取一题的选项，返回 (选项文本, 新下标)。"""
    if i >= end:
        return "", i
    L0 = lines[i]
    m_full = ABCD_ONE_LINE.search(L0)
    if m_full:
        return m_full.group(0).strip(), i + 1
    if OPTION_PREFIX.match(L0):
        parts: list[str] = []
        while i < end and OPTION_PREFIX.match(lines[i]):
            parts.append(lines[i])
            i += 1
        return "\n".join(parts).strip(), i
    if "\t" in L0 and i + 1 < end and "\t" in lines[i + 1]:
        parts12 = re.split(r"\s*\t\s*", L0) + re.split(r"\s*\t\s*", lines[i + 1])
        if len(parts12) >= 4:
            return (L0 + "\n" + lines[i + 1]).strip(), i + 2
    if i + 3 < end:
        quad = lines[i : i + 4]
        if all(len(x) < 120 for x in quad):
            return "\n".join(quad).strip(), i + 4
    if i < end:
        return lines[i].strip(), i + 1
    return "", i


def _starts_options(lines: list[str], i: int, end: int) -> bool:
    if i >= end:
        return False
    L = lines[i]
    if ABCD_ONE_LINE.search(L):
        return True
    if OPTION_PREFIX.match(L):
        return True
    if "\t" in L and i + 1 < end and "\t" in lines[i + 1]:
        parts12 = re.split(r"\s*\t\s*", L) + re.split(r"\s*\t\s*", lines[i + 1])
        if len(parts12) >= 4:
            return True
    if i + 3 < end and all(len(lines[i + k]) < 120 for k in range(4)):
        # 下一行可能是 end 边界外的「二、」大题标题，仍应用作「四行选项已结束」判断
        nxt = lines[i + 4] if (i + 4) < len(lines) else ""
        if nxt and (
            len(nxt) >= 40
            or SECTION_CHOICE.match(nxt)
            or try_question_start(nxt) is not None
        ):
            return True
    return False


def _parse_choice_implicit(lines: list[str], start: int, end: int, q_min: int, q_max: int) -> list[tuple[int, str]]:
    blocks: list[tuple[int, str]] = []
    i = start
    qn = q_min
    while i < end and qn <= q_max:
        if _is_noise_line(lines[i]):
            i += 1
            continue
        if SECTION_CHOICE.match(lines[i]) or SECTION_ANALYSIS.match(lines[i]):
            break
        head = try_question_start(lines[i])
        if head:
            qnx, t = head
            parts: list[str] = []
            if t:
                parts.append(t)
            i += 1
            while i < end:
                nxt = lines[i]
                if try_question_start(nxt):
                    break
                if SECTION_CHOICE.match(nxt) or SECTION_ANALYSIS.match(nxt):
                    break
                parts.append(nxt)
                i += 1
            text = "\n".join(parts).strip()
            if text and q_min <= qnx <= q_max:
                blocks.append((qnx, text))
            qn = max(qn, qnx + 1)
            continue
        stem_parts: list[str] = []
        while i < end:
            if _is_noise_line(lines[i]):
                i += 1
                continue
            if SECTION_CHOICE.match(lines[i]) or SECTION_ANALYSIS.match(lines[i]):
                break
            if try_question_start(lines[i]):
                break
            if _starts_options(lines, i, end):
                break
            stem_parts.append(lines[i])
            i += 1
        stem = "\n".join(stem_parts).strip()
        if stem_parts and i < end and try_question_start(lines[i]):
            continue
        if i < end and SECTION_CHOICE.match(lines[i]):
            break
        if not stem or len(stem) < 4:
            i += 1
            continue
        if "下列每小题" in stem or "多选、少选" in stem[:20]:
            continue
        if i >= end or not _starts_options(lines, i, end):
            if i < end:
                i += 1
            continue
        opts, i = _consume_options_block(lines, i, end)
        if not opts:
            continue
        blocks.append((qn, f"{stem}\n{opts}"))
        qn += 1
    return blocks


def _parse_analysis_implicit(lines: list[str], start: int, end: int) -> list[tuple[int, str]]:
    chunks: list[str] = []
    i = start
    while i < end:
        if _is_noise_line(lines[i]):
            i += 1
            continue
        if "结合材料" in lines[i] and "回答问题" in lines[i]:
            i += 1
            buf: list[str] = []
            while i < end:
                if "结合材料" in lines[i] and "回答问题" in lines[i]:
                    break
                if not _is_noise_line(lines[i]):
                    buf.append(lines[i])
                i += 1
            text = "\n".join(buf).strip()
            if text:
                chunks.append(text)
            continue
        i += 1
    out: list[tuple[int, str]] = []
    base = 34
    for j, ch in enumerate(chunks):
        out.append((base + j, ch))
    return out


def parse_exam_blocks(lines: list[str]) -> list[tuple[int, str]]:
    numbered = parse_numbered_blocks(lines)
    if len(numbered) >= 34:
        return sorted(numbered, key=lambda x: x[0])
    sec = _find_section_starts(lines)
    n = len(lines)
    i1 = sec.get("一", 0)
    i2 = sec.get("二", n)
    i3 = sec.get("三", n)
    if i1 >= n:
        return sorted(numbered, key=lambda x: x[0]) if numbered else []
    body1 = i1 + 1 if i1 < n else 0
    end1 = i2 if i2 < n else n
    end2 = i3 if i3 < n else n
    blocks = _parse_choice_implicit(lines, body1, end1, 1, 16)
    if i2 < n:
        body2 = i2 + 1
        blocks.extend(_parse_choice_implicit(lines, body2, end2, 17, 33))
    if i3 < n:
        body3 = i3 + 1
        blocks.extend(_parse_analysis_implicit(lines, body3, n))
    if not blocks and numbered:
        return sorted(numbered, key=lambda x: x[0])
    return sorted(blocks, key=lambda x: x[0])


def docx_to_items(path: Path, course_id: str) -> list[dict]:
    year = extract_year(path.name)
    lines = paragraphs_to_lines(Document(str(path)))
    blocks = parse_exam_blocks(lines)
    items: list[dict] = []
    for qn, text in blocks:
        items.append(
            {
                "id": f"kyzz-{year}-{qn:03d}",
                "courseId": course_id,
                "question": f"【{year}年考研政治】第{qn}题\n{text}",
                "answer": (
                    "（本条目由真题 Word 自动生成，未从「政治解析」PDF 抽取答案；"
                    "请对照解析册或教材补全，或自行编辑本 JSON。）"
                ),
                "tags": ["考研政治", year, "真题", path.stem],
            }
        )
    return items


def main() -> int:
    root = Path(__file__).resolve().parent.parent
    ap = argparse.ArgumentParser(description="考研政治真题 .docx → hongfan-bank.json")
    ap.add_argument(
        "--zhen-dir",
        type=Path,
        default=None,
        help="含 word 版真题的目录（默认自动查找 …/2009-2025政治真题）",
    )
    ap.add_argument(
        "--out",
        type=Path,
        default=root / "01历年真题及解析" / "hongfan-bank.json",
        help="输出 JSON 路径",
    )
    ap.add_argument(
        "--course",
        default="shizheng",
        help="courseId：与前端一致，默认 shizheng（考研政治综合）",
    )
    ap.add_argument(
        "--merge-existing",
        action="store_true",
        help="若输出文件已存在，合并其 items（按 id 去重，新题覆盖同 id）",
    )
    args = ap.parse_args()

    zhen = args.zhen_dir
    if zhen is None:
        zhen = find_politics_zhen_dir(root)
    if zhen is None or not zhen.is_dir():
        print("未找到「2009-2025政治真题」目录，请用 --zhen-dir 指定。", file=sys.stderr)
        return 1

    docx_files = sorted(zhen.rglob("*.docx"))
    if not docx_files:
        print(f"目录内无 .docx: {zhen}", file=sys.stderr)
        return 1

    all_items: list[dict] = []
    for f in docx_files:
        try:
            all_items.extend(docx_to_items(f, args.course))
        except Exception as e:  # noqa: BLE001
            print(f"跳过（解析失败）{f}: {e}", file=sys.stderr)

    if not all_items:
        print("未解析到任何题目。", file=sys.stderr)
        return 1

    out_path = args.out
    if not out_path.is_absolute():
        out_path = root / out_path

    payload: dict = {
        "version": 1,
        "sourceNote": "由 scripts/word_exam_to_hongfan_json.py 从考研政治真题 Word 生成；答案为占位，请合规补全。",
        "items": all_items,
    }

    if args.merge_existing and out_path.is_file():
        try:
            old = json.loads(out_path.read_text(encoding="utf-8"))
            old_items = old.get("items") if isinstance(old.get("items"), list) else []
            by_id = {str(x.get("id")): x for x in old_items if isinstance(x, dict) and x.get("id")}
            for it in all_items:
                by_id[str(it["id"])] = it
            payload["items"] = list(by_id.values())
            payload["sourceNote"] = (
                old.get("sourceNote", "") + " | 已合并 word_exam 导出。"
            )
        except (json.JSONDecodeError, OSError) as e:
            print(f"合并旧文件失败，将覆盖写入: {e}", file=sys.stderr)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"已写入 {out_path} ，共 {len(payload['items'])} 条（来自 {len(docx_files)} 个 docx）")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
